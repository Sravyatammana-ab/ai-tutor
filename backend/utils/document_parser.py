import os
import re
import hashlib
from typing import List, Dict, Tuple
from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    from utils.azure_ocr import AzureOCRService
except ImportError as e:
    AzureOCRService = None
    print(f"ERROR: Failed to import AzureOCRService: {e}")
    print("Please install: pip install azure-ai-documentintelligence azure-core")

try:
    from docx import Document
except ImportError:
    Document = None


def extract_text_from_docx(path: str) -> str:
    """Extract full text from DOCX using python-docx"""
    if not Document:
        raise ImportError("python-docx is not installed")
    
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return "\n".join(paragraphs)




def compute_file_hash(file_path: str) -> str:
    """Compute MD5 hash of file for duplicate detection"""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def detect_chapters_and_units(text: str) -> Dict:
    """
    Detect chapter and unit titles from text
    Returns dict with chapter_titles, unit_titles, and structured content
    """
    chapters = []
    units = []
    
    # Common patterns for chapters/units
    chapter_patterns = [
        r'(?i)^\s*(?:chapter|ch\.?|unit|lesson)\s*(\d+)[:\.]?\s*(.+?)$',
        r'(?i)^\s*(\d+)[:\.]\s*(.+?)$',  # Numbered headings
        r'(?i)^\s*([A-Z][A-Z\s]{3,})$',  # All caps headings
    ]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 3:
            continue
        
        # Check for chapter patterns
        for pattern in chapter_patterns:
            match = re.match(pattern, line)
            if match:
                if 'chapter' in pattern.lower() or 'ch.' in pattern.lower():
                    chapters.append({
                        'number': match.group(1) if match.lastindex >= 1 else str(len(chapters) + 1),
                        'title': match.group(2) if match.lastindex >= 2 else match.group(1),
                        'line_index': i
                    })
                elif 'unit' in pattern.lower() or 'lesson' in pattern.lower():
                    units.append({
                        'number': match.group(1) if match.lastindex >= 1 else str(len(units) + 1),
                        'title': match.group(2) if match.lastindex >= 2 else match.group(1),
                        'line_index': i
                    })
                break
    
    return {
        'chapters': chapters,
        'units': units,
        'chapter_count': len(chapters) if chapters else None,
        'unit_count': len(units) if units else None
    }


_azure_service = None


def get_azure_service():
    """
    Lazily create a single Azure Document Intelligence client.
    Prevents re-initializing heavy SDK objects per request.
    AZURE IS REQUIRED - will raise error if not available.
    """
    global _azure_service
    if _azure_service is not None:
        return _azure_service
    
    # Azure is REQUIRED - fail if not available
    if not AzureOCRService:
        raise ValueError(
            "AzureOCRService is not available. "
            "Please install: pip install azure-ai-documentintelligence azure-core"
        )
    
    # Check environment variables before attempting initialization
    from config import Config
    
    endpoint_set = bool(Config.AZURE_ENDPOINT and Config.AZURE_ENDPOINT.strip())
    key_set = bool(Config.AZURE_KEY and Config.AZURE_KEY.strip())
    
    if not endpoint_set or not key_set:
        error_msg = "Azure Document Intelligence is REQUIRED but not configured.\n"
        error_msg += f"  AZURE_ENDPOINT: {'SET' if endpoint_set else 'NOT SET or EMPTY'}\n"
        error_msg += f"  AZURE_KEY: {'SET' if key_set else 'NOT SET or EMPTY'}\n"
        error_msg += "Please set both AZURE_ENDPOINT and AZURE_KEY in your .env file or environment variables."
        raise ValueError(error_msg)
    
    try:
        _azure_service = AzureOCRService()
        print("✓ Azure Document Intelligence service initialized successfully")
    except Exception as e:
        import traceback
        error_msg = f"Failed to initialize Azure Document Intelligence: {str(e)}\n"
        error_msg += "Please verify:\n"
        error_msg += "  1. AZURE_ENDPOINT and AZURE_KEY are set correctly\n"
        error_msg += "  2. Values are valid (no extra spaces, correct format)\n"
        error_msg += "  3. Azure service is accessible"
        print(f"ERROR: {error_msg}")
        print(f"Traceback: {traceback.format_exc()}")
        raise ValueError(error_msg) from e
    
    return _azure_service


class DocumentParser:
    """Parse PDF and DOCX documents and chunk text using LangChain"""
    
    def __init__(self):
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?"]
        )
        
        # Azure Document Intelligence service - initialized only once via get_azure_service()
        self._azure_ocr_service_initialized = False
        self._azure_ocr_service = None
    
    def parse_document(self, file_path: str, file_ext: str) -> Tuple[str, Dict]:
        """
        Parse document and extract text content with metadata
        
        Returns:
            Tuple of (text_content, metadata)
        """
        # Compute file hash for duplicate detection
        file_hash = compute_file_hash(file_path)
        
        metadata = {
            'filename': os.path.basename(file_path),
            'file_type': file_ext,
            'file_size': os.path.getsize(file_path),
            'file_hash': file_hash,
            'source': 'azure_document_intelligence' if file_ext == 'pdf' else 'python-docx'
        }
        
        if file_ext == 'pdf':
            return self._parse_pdf(file_path, metadata)
        elif file_ext == 'docx':
            return self._parse_docx(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _parse_pdf(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Parse PDF using Azure Document Intelligence ONLY - no fallback"""
        # Initialize Azure service only once - will raise error if not available
        if not self._azure_ocr_service_initialized:
            print(f"[DocumentParser] Initializing Azure Document Intelligence service...")
            self._azure_ocr_service = get_azure_service()  # Will raise if Azure unavailable
            self._azure_ocr_service_initialized = True
        
        if not self._azure_ocr_service:
            raise ValueError(
                "Azure Document Intelligence service is not available. "
                "PDF extraction requires Azure OCR - no fallback available."
            )
        
        print(f"[DocumentParser] Using Azure Document Intelligence for PDF extraction...")
        try:
            # Extract text using Azure Document Intelligence - ONLY method
            text_content = self._azure_ocr_service.extract_text(file_path)
            
            if not text_content or not text_content.strip():
                raise ValueError(
                    "Azure Document Intelligence returned empty text. "
                    "The PDF may be empty or contain no readable text."
                )
            
            print(f"[DocumentParser] ✓ Azure Document Intelligence extraction successful ({len(text_content)} chars)")
            
            # Clean and normalize text
            text_content = self._clean_text(text_content)
            
            # Detect chapters and units
            structure_info = detect_chapters_and_units(text_content)
            metadata.update(structure_info)
            
            # Estimate page count based on text length
            word_count = len(text_content.split())
            estimated_pages = max(1, word_count // 500)
            metadata['page_count'] = estimated_pages
            metadata['source'] = 'azure_document_intelligence'
            
            return text_content, metadata
            
        except ValueError as e:
            # Re-raise ValueError as-is (already formatted)
            print(f"[DocumentParser] ✗ Azure Document Intelligence extraction failed: {str(e)}")
            raise
        except Exception as e:
            # Wrap any other exception with Azure error message
            error_msg = f"Azure OCR failed: {str(e)}"
            print(f"[DocumentParser] ✗ {error_msg}")
            raise ValueError(error_msg) from e
    
    def _parse_docx(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Parse DOCX file using python-docx"""
        try:
            text_content = extract_text_from_docx(file_path)
            
            if not text_content or not text_content.strip():
                raise ValueError("Failed to extract text from DOCX - document appears empty")
            
            # Clean and normalize text
            text_content = self._clean_text(text_content)
            
            # Detect chapters and units
            structure_info = detect_chapters_and_units(text_content)
            metadata.update(structure_info)
            
            # Count paragraphs
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
            metadata['paragraph_count'] = len(paragraphs)
            metadata['source'] = 'python-docx'
            word_count = len(text_content.split())
            metadata['page_count'] = max(1, word_count // 500) or 1
            
            return text_content, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text while keeping paragraph breaks"""
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Collapse tabs/multiple spaces but retain paragraph breaks
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove null chars
        text = text.replace('\x00', '')
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = None, chunk_overlap: int = None, metadata: Dict = None) -> List[Dict]:
        """
        Split text into chunks using LangChain RecursiveCharacterTextSplitter
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
            metadata: Document metadata for context
        
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if chunk_size is None:
            chunk_size = self.chunk_size
        if chunk_overlap is None:
            chunk_overlap = self.chunk_overlap
        
        # Update text splitter if parameters changed
        if chunk_size != self.chunk_size or chunk_overlap != self.chunk_overlap:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ".", "!", "?"]
            )
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Get chapter/unit info from metadata
        chapters = metadata.get('chapters', []) if metadata else []
        units = metadata.get('units', []) if metadata else []
        
        # Convert to list of dictionaries with enhanced metadata
        chunk_list = []
        lines = text.split('\n')
        current_chapter = None
        current_unit = None
        
        for idx, chunk_text in enumerate(chunks):
            if not chunk_text.strip():
                continue
            
            # Find which chapter/unit this chunk belongs to
            chunk_start_pos = text.find(chunk_text)
            chunk_line_index = text[:chunk_start_pos].count('\n') if chunk_start_pos >= 0 else 0
            
            # Determine chapter and unit for this chunk
            for chapter in chapters:
                if chapter['line_index'] <= chunk_line_index:
                    current_chapter = chapter
                else:
                    break
            
            for unit in units:
                if unit['line_index'] <= chunk_line_index:
                    current_unit = unit
                else:
                    break
            
            chunk_metadata = {
                    'text': chunk_text,
                    'chunk_index': idx,
                    'length': len(chunk_text)
            }
            
            # Add chapter/unit metadata if available
            if current_chapter:
                chunk_metadata['chapter_number'] = current_chapter.get('number')
                chunk_metadata['chapter_title'] = current_chapter.get('title')
            
            if current_unit:
                chunk_metadata['unit_number'] = current_unit.get('number')
                chunk_metadata['unit_title'] = current_unit.get('title')
            
            chunk_list.append(chunk_metadata)
        
        return chunk_list
